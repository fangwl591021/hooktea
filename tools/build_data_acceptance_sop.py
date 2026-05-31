from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"D:\OneDrive\文件\New project 6\docs\HookTea 管理站-資料層每日驗收SOP.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_table_borders(table, color="D8DEE9"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style = doc.styles["Heading {}".format(level)]
    p.style = style
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc, title, body, fill="F4F6F9", color="0B2545"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    apply_table_geometry(table, [9360])
    set_table_borders(table, "D8DEE9")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Microsoft JhengHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Microsoft JhengHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string("334155")
    doc.add_paragraph()


def add_status_table(doc):
    add_heading(doc, "資料層現況", 1)
    rows = [
        ["資料類型", "目前來源", "Wasabi 狀態", "驗收重點"],
        ["低風險：課程、商城、影音", "可切 Wasabi 優先，保留 R2/KV fallback", "快照、一致性驗證、雙寫觀察", "筆數與 hash 一致"],
        ["高風險：會員、點數、點數總表、訂單", "R2 live 優先，KV fallback", "快照與雙寫觀察，尚未作為主讀取來源", "贈扣點不受 KV 滿額阻擋"],
        ["KV", "兼容與回滾用途", "不是唯一寫入承載", "KV 滿額時不能讓營運操作失敗"],
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    apply_table_geometry(table, [2050, 2500, 2450, 2360])
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF5")
                set_cell_text(cell, text, bold=True, color="1F3A5F")
            else:
                set_cell_text(cell, text)


def add_daily_steps(doc):
    add_heading(doc, "每日驗收流程", 1)
    steps = [
        "進入後台「營運工具 > Wasabi 遷移檢查」。",
        "先按「執行檢查」，確認連線健康、讀取來源狀態、雙寫觀察狀態都有更新。",
        "按「每日總檢查」。系統會一次執行 Wasabi 健康檢查、低風險一致性、高風險一致性。",
        "總檢查顯示通過時，截圖留存即可作為當日驗收紀錄。",
        "若任一區塊顯示失敗，依本文件「異常判讀」處理，不要直接切換主庫。",
    ]
    for item in steps:
        add_number(doc, item)


def add_acceptance_table(doc):
    add_heading(doc, "通過標準", 1)
    rows = [
        ["檢查項目", "畫面應顯示", "通過判斷"],
        ["Wasabi 連線健康", "PutObject / HeadObject / GetObject / DeleteObject 通過", "四項皆通過"],
        ["低風險一致性", "課程、商城、影音皆為一致", "本機筆數與 Wasabi 筆數一致，hash 一致"],
        ["高風險一致性", "會員、點數、點數總表、訂單皆為一致", "本機筆數與 Wasabi 筆數一致，hash 一致"],
        ["讀取來源狀態", "高風險：R2 live 優先 / KV fallback", "顯示清楚，避免誤認已切 Wasabi 主庫"],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    apply_table_geometry(table, [2600, 3900, 2860])
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF5")
                set_cell_text(cell, text, bold=True, color="1F3A5F")
            else:
                set_cell_text(cell, text)


def add_point_test(doc):
    add_heading(doc, "贈點測試流程", 1)
    add_callout(
        doc,
        "測試目的",
        "確認 KV 當日寫入額度已滿時，點數贈扣仍能由 R2 live 接住，並能進入 Wasabi 高風險快照驗證。",
        fill="FFF7ED",
        color="7A5A00",
    )
    for item in [
        "在會員 CRM 開啟一位測試會員。",
        "執行「贈點」並輸入小額點數，例如 1 點或 10 點，原因填寫 test。",
        "確認畫面沒有出現 KV put() limit exceeded for the day。",
        "回到 Wasabi 遷移檢查，按「執行檢查」。",
        "確認「高風險雙寫觀察」中的會員點數與點數進出總表為最後寫入成功。",
        "按「驗證高風險一致」或「每日總檢查」，確認高風險一致性通過。",
    ]:
        add_number(doc, item)


def add_failure_guide(doc):
    add_heading(doc, "異常判讀", 1)
    rows = [
        ["異常畫面", "可能原因", "處理方式"],
        ["KV put() limit exceeded", "KV 今日寫入額度已滿，或仍有未改成 R2 live 的寫入點", "先確認是否已部署最新版；若仍發生，記錄操作位置與錯誤畫面再修補該寫入點"],
        ["低風險不一致", "課程、商城或影音資料有新異動但尚未重新快照", "先按匯出低風險快照，再按驗證；若仍不一致，檢查商品/課程數量是否符合預期"],
        ["高風險不一致", "會員、點數或訂單剛異動，快照尚未更新或同步失敗", "先按匯出高風險快照，再按驗證；若仍不一致，查看紅色項目是哪一組資料"],
        ["Wasabi 健康失敗", "Wasabi Secret、Bucket、Region 或網路連線異常", "先不要切換來源；檢查 Worker Secrets 與 Wasabi 服務狀態"],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    apply_table_geometry(table, [2500, 3300, 3560])
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_shading(cell, "FEE2E2")
                set_cell_text(cell, text, bold=True, color="9B1C1C")
            else:
                set_cell_text(cell, text)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Title", 22, "0B2545", 0, 10),
        ("Subtitle", 11, "64748B", 0, 12),
        ("Heading 1", 16, "1F3A5F", 14, 6),
        ("Heading 2", 13, "2E74B5", 10, 4),
        ("Heading 3", 12, "1F4D78", 8, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("HookTea 管理站 資料層驗收 SOP")
    r.font.name = "Microsoft JhengHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("64748B")


def build():
    doc = Document()
    setup_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("HookTea 管理站 資料層每日驗收 SOP")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("適用範圍：Wasabi 遷移檢查、R2 live fallback、低風險與高風險資料一致性驗收")

    add_callout(
        doc,
        "目前結論",
        "低風險資料已可切 Wasabi 優先讀取；高風險資料目前採 R2 live 優先、KV fallback，Wasabi 先作為快照與雙寫觀察，不作為會員、點數、訂單的主讀取庫。",
        fill="ECFDF5",
        color="047857",
    )

    add_status_table(doc)
    add_daily_steps(doc)
    add_acceptance_table(doc)
    add_point_test(doc)
    add_failure_guide(doc)

    add_heading(doc, "交付前提醒", 1)
    for item in [
        "每日總檢查通過，代表當下快照與來源一致，不代表可以立即移除 KV。",
        "高風險資料仍須觀察一段真實營運寫入，再討論是否切 Wasabi 主讀取。",
        "任何紅色失敗項目都應先截圖與記錄時間，再進行修補或重新匯出。",
        "業主驗收時，建議同時保留每日總檢查截圖與一次贈點測試截圖。",
    ]:
        add_bullet(doc, item)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
